"""
Generate a navigational audit Word report (.docx) from inspection data.
Matches the structure of Yatendra's Excel-based audit workbook output.
"""

from __future__ import annotations

import io
from datetime import datetime
from typing import Any, Dict, List, Optional

import matplotlib
matplotlib.use("Agg")  # headless — no display needed
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.services.vhi import compute_audit_score


# ─── helpers ────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    """Set cell background shading."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_border(cell, **borders):
    """Set borders on a table cell. borders keys: top/bottom/left/right."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border_data = borders.get(side, {"val": "single", "sz": "4", "color": "AAAAAA"})
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), border_data.get("val", "single"))
        el.set(qn("w:sz"), border_data.get("sz", "4"))
        el.set(qn("w:color"), border_data.get("color", "AAAAAA"))
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _bold_cell(cell, text: str, font_size: int = 10, color: str = None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(font_size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _score_to_color(pct: float) -> str:
    """Return a hex color based on percentage score."""
    if pct >= 80:
        return "92D050"  # green
    if pct >= 65:
        return "FFFF00"  # yellow
    if pct >= 50:
        return "FFC000"  # amber
    if pct >= 35:
        return "FF0000"  # red
    return "C00000"     # dark red


# ─── chart generation ───────────────────────────────────────────────────────

def _bar_chart_png(
    categories: List[str],
    percentages: List[float],
    title: str,
    width_in: float = 5.5,
    height_in: float = 3.0,
) -> bytes:
    """Render a horizontal bar chart and return PNG bytes."""
    n = len(categories)
    fig, ax = plt.subplots(figsize=(width_in, max(height_in, n * 0.45 + 0.8)))

    colors = ["#92D050" if p >= 80 else "#FFC000" if p >= 65 else "#FF0000" for p in percentages]
    y_pos = np.arange(n)
    bars = ax.barh(y_pos, percentages, color=colors, height=0.55, edgecolor="white")

    # Value labels
    for bar, val in zip(bars, percentages):
        ax.text(
            min(val + 1.5, 102),
            bar.get_y() + bar.get_height() / 2,
            f"{val:.1f}%",
            va="center",
            ha="left",
            fontsize=9,
            fontweight="bold",
        )

    ax.set_yticks(y_pos)
    ax.set_yticklabels(categories, fontsize=9)
    ax.set_xlim(0, 115)
    ax.set_xlabel("Percentage Score (%)", fontsize=9)
    ax.set_title(title, fontsize=11, fontweight="bold", pad=10)
    ax.axvline(x=80, color="green", linestyle="--", linewidth=0.8, alpha=0.6, label="≥80% (A)")
    ax.axvline(x=65, color="gold", linestyle="--", linewidth=0.8, alpha=0.6, label="≥65% (B)")
    ax.legend(fontsize=7, loc="lower right")
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.grid(axis="x", alpha=0.3)
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150, bbox_inches="tight")
    plt.close(fig)
    buf.seek(0)
    return buf.read()


# ─── main report function ────────────────────────────────────────────────────

def generate_report(
    inspection: Dict[str, Any],
    scores: List[Dict[str, Any]],
    vessel_name: Optional[str] = None,
    company_name: Optional[str] = None,
) -> bytes:
    """
    Build the Word report and return raw .docx bytes.

    Parameters
    ----------
    inspection : dict   — from inspection_to_dict()
    scores     : list   — from GET /scores endpoint (includes item_name, guidance_note etc.)
    vessel_name: str    — optional override (falls back to "—")
    company_name: str   — optional override
    """
    doc = Document()

    # ── page margins ──────────────────────────────────────────────────────
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # ── compute audit score ───────────────────────────────────────────────
    raw_scores = [
        {
            "score": s["score"],
            "category": s["category"],
            "assessment_type": s.get("assessment_type", "static"),
            "weight": s.get("weight", 1),
        }
        for s in scores
    ]
    result = compute_audit_score(raw_scores)
    pct = result["percentage"]
    avg = result["average"]
    grade = _grade(pct)

    cat_summary = result["category_summary"]
    static_cats = [c for c in cat_summary if c["assessment_type"] == "static"]
    dynamic_cats = [c for c in cat_summary if c["assessment_type"] == "dynamic"]

    # ── Cover / Header ────────────────────────────────────────────────────
    _heading(doc, "NAVIGATIONAL AUDIT REPORT", 0)
    _heading(doc, "Nivyash Marine Management Services", 1)

    doc.add_paragraph()

    _info_table(doc, [
        ("Vessel Name",      vessel_name or "—"),
        ("Company",          company_name or "Nivyash Marine Management Services"),
        ("Analysis Period",  _fmt_date(inspection.get("inspection_date"))),
        ("Port / Zone",      inspection.get("port") or "—"),
        ("Audit Date",       _fmt_date(inspection.get("inspection_date"))),
        ("Report Generated", datetime.utcnow().strftime("%d %b %Y")),
        ("Overall Grade",    f"{grade}  ({pct:.1f}%)"),
    ])

    doc.add_paragraph()

    # ── 1. Key Indicators ─────────────────────────────────────────────────
    _section_title(doc, "1. Key Indicators")

    ki_table = doc.add_table(rows=1, cols=2)
    ki_table.style = "Table Grid"
    ki_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = ki_table.rows[0].cells
    _bold_cell(hdr[0], "Indicator", color="FFFFFF")
    _bold_cell(hdr[1], "Value", color="FFFFFF")
    _set_cell_bg(hdr[0], "1F497D")
    _set_cell_bg(hdr[1], "1F497D")

    ki_rows = [
        ("Total Checklist Items",    str(result["total_subjects"])),
        ("Total Subjects Assessed",  str(result["total_assessed"])),
        ("Total Not Sighted (NS)",   str(result["total_ns"])),
        ("Maximum Possible Score",   str(result["max_score"])),
        ("Total Score Obtained",     str(result["total_score"])),
        ("Average Score (out of 5)", f"{avg:.2f}"),
        ("Percentage Score",         f"{pct:.2f}%"),
        ("Overall Grade",            grade),
    ]
    for i, (label, value) in enumerate(ki_rows):
        row = ki_table.add_row().cells
        row[0].text = label
        row[0].paragraphs[0].runs[0].font.size = Pt(10)
        p = row[1].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(value)
        run.bold = (label in ("Percentage Score", "Overall Grade"))
        run.font.size = Pt(10)
        if label == "Overall Grade":
            _set_cell_bg(row[1], _score_to_color(pct))
        elif i % 2 == 0:
            _set_cell_bg(row[0], "EBF3FB")
            _set_cell_bg(row[1], "EBF3FB")

    doc.add_paragraph()

    # ── 2. Static Assessment Aspects ──────────────────────────────────────
    _section_title(doc, "2. Static Assessment — Aspect-wise Performance")
    _category_table(doc, static_cats)
    doc.add_paragraph()

    # ── 3. Dynamic Assessment Aspects ─────────────────────────────────────
    _section_title(doc, "3. Dynamic Assessment — Aspect-wise Performance")
    _category_table(doc, dynamic_cats)
    doc.add_paragraph()

    # ── 4. Performance Charts ─────────────────────────────────────────────
    _section_title(doc, "4. Performance Charts")

    if static_cats:
        png = _bar_chart_png(
            [c["category"] for c in static_cats],
            [c["percentage_score"] for c in static_cats],
            "Static Assessment — Category Performance",
        )
        buf = io.BytesIO(png)
        doc.add_picture(buf, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    if dynamic_cats:
        png = _bar_chart_png(
            [c["category"] for c in dynamic_cats],
            [c["percentage_score"] for c in dynamic_cats],
            "Dynamic Assessment — Category Performance",
        )
        buf = io.BytesIO(png)
        doc.add_picture(buf, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

    # ── 5. Deficiency Summary ─────────────────────────────────────────────
    deficiencies = [s for s in scores if s.get("is_deficiency")]
    if deficiencies:
        _section_title(doc, "5. Deficiency Summary")
        def_table = doc.add_table(rows=1, cols=4)
        def_table.style = "Table Grid"
        def_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdrs = def_table.rows[0].cells
        for cell, label in zip(hdrs, ["#", "Item", "Score", "Comment"]):
            _bold_cell(cell, label, color="FFFFFF")
            _set_cell_bg(cell, "C00000")
        def_table.columns[0].width = Inches(0.3)
        def_table.columns[1].width = Inches(3.5)
        def_table.columns[2].width = Inches(0.6)
        def_table.columns[3].width = Inches(2.0)

        for i, s in enumerate(deficiencies, 1):
            row = def_table.add_row().cells
            row[0].text = str(i)
            row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            row[1].text = s.get("item_name", "")
            p = row[2].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(s["score"]))
            run.bold = True
            run.font.color.rgb = RGBColor.from_string("C00000")
            row[3].text = s.get("comment") or ""
            if i % 2 == 0:
                for cell in row:
                    _set_cell_bg(cell, "FFF2CC")
        doc.add_paragraph()

    # ── 6. Complete Scores ────────────────────────────────────────────────
    _section_title(doc, "6. Complete Score Sheet")

    # Group by assessment_type → category
    from collections import defaultdict
    groups: Dict[tuple, list] = defaultdict(list)
    for s in scores:
        key = (s.get("assessment_type", "static"), s["category"])
        groups[key].append(s)

    for (atype, cat), items in sorted(groups.items()):
        _subheading(doc, f"{atype.title()} — {cat}")
        score_table = doc.add_table(rows=1, cols=4)
        score_table.style = "Table Grid"
        hdrs2 = score_table.rows[0].cells
        for cell, label in zip(hdrs2, ["Code/Item", "Item Description", "Score", "Comment"]):
            _bold_cell(cell, label, color="FFFFFF")
            _set_cell_bg(cell, "1F497D")
        score_table.columns[0].width = Inches(0.8)
        score_table.columns[1].width = Inches(3.5)
        score_table.columns[2].width = Inches(0.6)
        score_table.columns[3].width = Inches(1.5)

        for i, s in enumerate(items):
            row = score_table.add_row().cells
            row[0].text = s.get("checklist_item_id", "")[:8] + "…"
            row[1].text = s.get("item_name", "")
            row[1].paragraphs[0].runs[0].font.size = Pt(9)

            p = row[2].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sc = s.get("score")
            run = p.add_run(str(sc) if sc is not None else "—")
            run.font.size = Pt(9)
            if isinstance(sc, int) and sc < 3:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string("C00000")
            elif sc == "NS":
                run.font.color.rgb = RGBColor.from_string("888888")

            row[3].text = s.get("comment") or ""
            row[3].paragraphs[0].runs[0].font.size = Pt(9) if row[3].paragraphs[0].runs else Pt(9)
            if i % 2 == 0:
                for cell in row:
                    _set_cell_bg(cell, "F5F8FB")
        doc.add_paragraph()

    # ── Footer note ───────────────────────────────────────────────────────
    p = doc.add_paragraph(
        "This report has been auto-generated from the MarinePulse audit platform. "
        "Executive summary, narrative findings, and recommendations should be added manually."
    )
    p.style = "Normal"
    p.runs[0].font.size = Pt(8)
    p.runs[0].font.color.rgb = RGBColor.from_string("888888")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Return bytes ──────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ─── small doc helpers ───────────────────────────────────────────────────────

def _grade(pct: float) -> str:
    if pct >= 80: return "A"
    if pct >= 65: return "B"
    if pct >= 50: return "C"
    if pct >= 35: return "D"
    return "F"


def _fmt_date(iso: Optional[str]) -> str:
    if not iso:
        return "—"
    try:
        dt = datetime.fromisoformat(iso.replace("Z", "+00:00"))
        return dt.strftime("%d %b %Y")
    except Exception:
        return iso


def _heading(doc, text: str, level: int):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _section_title(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor.from_string("1F497D")
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def _subheading(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor.from_string("375623")
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)


def _info_table(doc, rows: List[tuple]):
    """Two-column key/value info table."""
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for label, value in rows:
        row = tbl.add_row().cells
        row[0].text = label
        row[0].paragraphs[0].runs[0].bold = True
        row[0].paragraphs[0].runs[0].font.size = Pt(10)
        _set_cell_bg(row[0], "D9E1F2")
        row[1].text = value
        row[1].paragraphs[0].runs[0].font.size = Pt(10)


def _category_table(doc, cats: List[Dict[str, Any]]):
    if not cats:
        doc.add_paragraph("No data available.")
        return
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdrs = tbl.rows[0].cells
    for cell, label in zip(hdrs, ["Category", "Subjects Assessed", "Average Score (/ 5)", "Percentage Score"]):
        _bold_cell(cell, label, color="FFFFFF")
        _set_cell_bg(cell, "1F497D")

    for i, c in enumerate(cats):
        row = tbl.add_row().cells
        row[0].text = c["category"]
        row[0].paragraphs[0].runs[0].font.size = Pt(10)

        for j, val in enumerate([
            str(c["subjects_assessed"]),
            f"{c['average_score']:.2f}",
            f"{c['percentage_score']:.1f}%",
        ], 1):
            p = row[j].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            run.font.size = Pt(10)
            if j == 3:
                run.bold = True
                _set_cell_bg(row[j], _score_to_color(c["percentage_score"]))

        if i % 2 == 0:
            _set_cell_bg(row[0], "EBF3FB")
            _set_cell_bg(row[1], "EBF3FB")
            _set_cell_bg(row[2], "EBF3FB")
